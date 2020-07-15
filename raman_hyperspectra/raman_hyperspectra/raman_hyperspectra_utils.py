# raman_hyperspectra_utils
#
#toolbox of utilities
#
# create_report : generates a word report
# mode_selection : 
# pickle
# save_image
# select_file
#
#3rd party dependencies : matplotlib, docx, pickle




def save_image(func, fig) -> None:

    '''
    save the plots/images generated by the funtion func
    Arguments:
        func (string): name of the function which generates the graphs
    Returns:
        None
    '''

    import time
    import matplotlib.pyplot as plt
    fig.savefig(func + " " + time.strftime("%H_%M_%S")+".png")

def select_file(path,file_extension = ("txt files","*.txt"), title = "Select file"):

    '''
    interactive selection of a file
    Arguments:
        path (string): path of the directory conataining the file
        file_extension (string): filter of file_extension
    Returns:
        file_raw (string): sected file
    '''

    from tkinter import filedialog
    import tkinter as tk

    root = tk.Tk()
    file_raw =  filedialog.askopenfilename(initialdir = path,
                     title = title, filetypes = (file_extension,("all files","*.*")))
    root.destroy()
    return file_raw

def create_report(file_info, file_init, sheet, path_image, da):

    from docx import Document
    from docx.shared import Cm #Inches
    import numpy as np
    import re
    import os

    name_dir  = os.getcwd()
    file_png = os.listdir()

    document = Document()

    motif=re.compile("\d+_\d+_\d+")
    date_exploitation = re.findall(motif, name_dir)[0]
    name_hyperspectra = name_dir.split('\\')[-1][0:-len(date_exploitation)]

    document.add_heading("Traitement de l'hyperspectre "+name_hyperspectra , 0)

    document.add_heading('Informations générales ', level=1)
    p = document.add_paragraph("date d'exploitation :"+"- ".join(map(str, date_exploitation.strip().split("_"))))

    dic_info = read_RAMAN_WITEC_information(file_info)
    for key, value in dic_info.items():
        p = document.add_paragraph(f'{key} : {value}')

    p = document.add_paragraph(f'Spectral resolution[cm-1] : {round(np.diff(da.lbd.values).mean(),2)}')

    *_  ,label, _ = init_model(file_init, sheet)
    p = document.add_paragraph('fit model')
    for X in label:
        p = document.add_paragraph(X)

    document.add_heading("Image de l'échantillon", level=1)
    document.add_picture(path_image, width=Cm(17))

    document.add_heading('Analyse du bruit', level=1)
    for x in [x for x in file_png if 'plot_histo' in x]:
        document.add_picture(x, width=Cm(17))

    document.add_heading('Analyse de Lbd_0', level=1)
    for x in [x for x in file_png if 'plot_noises' in x]:
        document.add_picture(x, width=Cm(17))

    document.add_heading('Analyse du R2', level=1)
    for x in [x for x in file_png if 'plot_R2' in x]:
        document.add_picture(x, width=Cm(17))
    for x in [x for x in file_png if 'plot_lbd_0' in x]:
        document.add_picture(x, width=Cm(17))

    document.add_heading('Analyse des maximums', level=1)
    for x in [x for x in file_png if 'Stat_MAX' in x]:
        document.add_picture(x, width=Cm(17))

    document.add_heading('Analyse des phases', level=1)
    for x in [x for x in file_png if 'phase_maximums' in x]:
        document.add_picture(x, width=Cm(17))
    for x in [x for x in file_png if 'trace_phase_menu' in x]:
        document.add_picture(x, width=Cm(17))

    document.save('demo.docx')

def pickle(path_save, filename, mode, *args):

    '''
    picking of of variables
    Arguments:
        path_save (string): path of the directory where to store the pickle file
        filemane (string): name of the pickle file (witout extension)
        mode (string): 'write' to write a new file
                       'append' to append
                       'read' to read the pickle
        *args (lis): variables to pickle
    Returns:
        a pickle file store in the directory located at path_save
    '''

    import pickle
    import os

    data = []

    if mode == 'write':
        with open(os.path.join(path_save , filename + ".pkl"), "wb") as pickle_file:
            pickle.dump(filename,pickle_file)
            for item in args:
                pickle.dump(item,pickle_file)

    elif mode == 'append':
        with open(os.path.join(path_save , filename + ".pkl"), "ab") as pickle_file:
            for item in args:
                pickle.dump(item,pickle_file)

    elif mode == 'read':
        with open(os.path.join(path_save , filename ), "rb") as pickle_file:
            try:
                while True:
                    data.append(pickle.load(pickle_file))
            except EOFError:
                pass
    else:
        raise ValueError('invalid mode must be write, append or read')

    return data

def mode_selection() :
    '''
    selection of operating option
    
    Argument:
    
    Reurns:
        choice_file (string): spectrum/interp spectrum/full spectrum/interp full spectrum
        choice_method (list): max Stat, Phase imaging, PCA, NMF, robust PCA
        choice_shift(string): Rayleigh/Ramn peak/no correction
    '''
    import tkinter as tk
    global choice_file
    global choice_shift
    
    root = tk.Tk()

    file = [
        ("spectrum", 1),
        ("interp spectrum", 2),
        ("full spectrum", 3),
        ("interp full spectrum", 4),
    ]
    
    shift = [
        ("Rayleigh", 1),
        ("Raman peak", 2),
        ("no correction", 3),
    ]
    
    choice_file = "spectrum"
    def Choice_File(text, v):
        global choice_file
        choice_file = text
        
    choice_shift = "Rayleigh"
    def Choice_Shift(text, v):
        global choice_shift
        choice_shift = text

    choice_method = []
    def Choice_method(text,v):
        if text in choice_method:
            choice_method.remove(text)
        else:
            choice_method.append(text)

    #
    # choice of the working spectrum
    #
    varfile = tk.IntVar()
    varfile.set(file[0][1])

    tk.Label(root, text='Choose the working file:').pack(anchor = tk.W)

    for txt, val in file:
        tk.Radiobutton(root, text = txt, variable = varfile, value=val,
            command=lambda t = txt, v = varfile: Choice_File(t, v)).pack(anchor=tk.NW)
    
    #
    # choice of the wavelength correction
    #
    varshift = tk.IntVar()
    varshift.set(shift[0][1])
    tk.Label(root, text='Choose the wavelength correction:').pack(anchor = tk.W)
    for txt, val in shift:
        tk.Radiobutton(root, text = txt, variable = varshift, value=val,
            command=lambda t = txt, v = varshift: Choice_Shift(t, v)).pack(anchor=tk.NW)
    
    #
    # choice of the method(s)
    #
    var_max_Stat = tk.IntVar()
    var_phase_imaging = tk.IntVar()
    var_PCA = tk.IntVar()
    var_NMF = tk.IntVar()
    var_robust_PCA = tk.IntVar()


    tk.Label(root, text='Choose method:').pack(anchor = tk.W)

    tk.Checkbutton(root, text = "max Stat", variable = var_max_Stat, onvalue = 1, offvalue = 0,
            command=lambda t = "max Stat", v = var_max_Stat: Choice_method(t,v)).pack(anchor=tk.NW)
    tk.Checkbutton(root, text = "Phase imaging", variable = var_phase_imaging, onvalue = 1, offvalue =0,
            command=lambda t = "Phase imaging", v = var_phase_imaging: Choice_method(t,v)).pack(anchor=tk.NW)
    tk.Checkbutton(root, text = "PCA", variable = var_PCA, onvalue = 1, offvalue =0,
            command=lambda t = "PCA", v=var_PCA: Choice_method(t,v)).pack(anchor=tk.NW)
    tk.Checkbutton(root, text = "NMF", variable =var_NMF, onvalue=1, offvalue = 0,
            command=lambda t = "NMF", v = var_NMF: Choice_method(t,v)).pack(anchor=tk.NW)
    tk.Checkbutton(root, text = "robust PCA", variable = var_robust_PCA, onvalue = 1, offvalue = 0,
            command=lambda t = "robust PCA", v = var_robust_PCA: Choice_method(t,v)).pack(anchor=tk.NW)

    root.mainloop()
    
    return choice_file, choice_method, choice_shift